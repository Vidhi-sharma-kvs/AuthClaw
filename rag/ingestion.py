import io
import os
import pypdf
import docx

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Safely extracts text content from PDF, DOCX, TXT, or MD files.
    """
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == ".pdf":
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = pypdf.PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Error parsing PDF file: {str(e)}")
            
    elif ext == ".docx":
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            text_parts = []
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text for cell in row.cells if cell.text]
                    if row_cells:
                        text_parts.append(" | ".join(row_cells))
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Error parsing DOCX file: {str(e)}")
            
    elif ext in (".txt", ".md", ".markdown"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(f"Error parsing text file: {str(e)}")
            
    else:
        # Fallback to text decoding
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            raise ValueError(f"Unsupported file format: {ext}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """
    Splits text into chunks of specified character length with overlap.
    """
    chunks = []
    text = text.strip()
    if not text:
        return chunks
        
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
            
        if end >= text_len:
            break
            
        start += (chunk_size - overlap)
        # Prevent infinite loops in case configuration is broken
        if chunk_size - overlap <= 0:
            start += chunk_size
            
    return chunks
