import io
import csv
import os
import zipfile
import logging
import defusedxml.ElementTree as ET
from dataclasses import dataclass
from typing import Dict, List, Any, Optional

logger = logging.getLogger("authclaw.document_processing.parsers")


@dataclass
class ExtractedDocumentPage:
    page_number: int
    text: str
    source: str = "native"


@dataclass
class DocumentExtractionResult:
    text: str
    pages: List[ExtractedDocumentPage]
    extraction_method: str
    ocr_status: str = "not_required"
    ocr_error: Optional[str] = None

def parse_txt_or_md(file_bytes: bytes) -> str:
    """Parses TXT or MD files."""
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Error parsing TXT/MD: {e}")
        return ""

def parse_csv(file_bytes: bytes) -> str:
    """Parses CSV content and formats it as a Markdown table."""
    try:
        text_stream = io.StringIO(file_bytes.decode("utf-8", errors="ignore"))
        reader = csv.reader(text_stream)
        rows = list(reader)
        if not rows:
            return ""
        
        # Format as markdown table
        markdown_parts = []
        for idx, row in enumerate(rows):
            # Clean cells
            row_cells = [cell.strip().replace("|", "\\|") for cell in row]
            markdown_parts.append(f"| {' | '.join(row_cells)} |")
            if idx == 0:
                # Add separator
                markdown_parts.append(f"| {' | '.join(['---'] * len(row_cells))} |")
        return "\n".join(markdown_parts)
    except Exception as e:
        logger.error(f"Error parsing CSV: {e}")
        return ""

def parse_xlsx(file_bytes: bytes) -> str:
    """
    Parses XLSX worksheets natively using Python zipfile and xml parser.
    Avoids openpyxl dependency for sandboxed/offline reliability.
    Formats the cells as a Markdown table.
    """
    try:
        file_like = io.BytesIO(file_bytes)
        if not zipfile.is_zipfile(file_like):
            return "Invalid Excel file format (not a valid zip)."
            
        with zipfile.ZipFile(file_like) as z:
            # 1. Load Shared Strings (if exist)
            shared_strings = []
            if "xl/sharedStrings.xml" in z.namelist():
                ss_data = z.read("xl/sharedStrings.xml")
                root = ET.fromstring(ss_data)
                # Namespace for spreadsheetml
                ns = {'ns': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
                for t in root.findall('.//ns:t', ns):
                    shared_strings.append(t.text or "")
                    
            # 2. Find and parse Sheet1 (or multiple sheets)
            sheet_files = [name for name in z.namelist() if name.startswith("xl/worksheets/sheet")]
            if not sheet_files:
                return "No worksheets found in Excel file."
                
            all_sheets_content = []
            ns = {'ns': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
            
            for sheet_file in sorted(sheet_files):
                sheet_data = z.read(sheet_file)
                root = ET.fromstring(sheet_data)
                
                # Extract grid coordinates and cells
                rows_dict: Dict[int, Dict[int, str]] = {}
                for row_el in root.findall('.//ns:row', ns):
                    row_idx = int(row_el.attrib.get("r", 1))
                    rows_dict[row_idx] = {}
                    
                    for c_el in row_el.findall('ns:c', ns):
                        # Cell coordinates like A1, B2
                        cell_ref = c_el.attrib.get("r", "")
                        # Parse column index from reference (A=1, B=2, etc.)
                        col_letters = "".join(filter(str.isalpha, cell_ref))
                        col_idx = 0
                        for char in col_letters:
                            col_idx = col_idx * 26 + (ord(char.upper()) - 64)
                            
                        val = ""
                        v_el = c_el.find('ns:v', ns)
                        if v_el is not None:
                            val = v_el.text or ""
                            # If type is shared string
                            if c_el.attrib.get("t") == "s":
                                try:
                                    val_idx = int(val)
                                    if 0 <= val_idx < len(shared_strings):
                                        val = shared_strings[val_idx]
                                except ValueError:
                                    pass
                        rows_dict[row_idx][col_idx] = val
                
                if not rows_dict:
                    continue
                    
                # Format sheet rows as Markdown Table
                max_row = max(rows_dict.keys())
                # Find maximum column across all rows
                max_col = 0
                for r in rows_dict.values():
                    if r.keys():
                        max_col = max(max_col, max(r.keys()))
                        
                if max_col == 0:
                    continue
                    
                markdown_lines = []
                markdown_lines.append(f"### Sheet: {os.path.basename(sheet_file).replace('.xml', '')}")
                
                # Loop through all rows in grid
                header_written = False
                for r_idx in range(1, max_row + 1):
                    row_cells = []
                    for c_idx in range(1, max_col + 1):
                        cell_val = rows_dict.get(r_idx, {}).get(c_idx, "").strip()
                        row_cells.append(cell_val.replace("|", "\\|"))
                    
                    markdown_lines.append(f"| {' | '.join(row_cells)} |")
                    if r_idx == 1 and not header_written:
                        markdown_lines.append(f"| {' | '.join(['---'] * len(row_cells))} |")
                        header_written = True
                all_sheets_content.append("\n".join(markdown_lines))
                
            return "\n\n".join(all_sheets_content)
    except Exception as e:
        logger.error(f"Excel zip/xml parsing failed: {e}")
        return ""

def parse_docx(file_bytes: bytes) -> str:
    """Parses DOCX text and tables using python-docx."""
    try:
        import docx
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells if cell.text]
                if cells:
                    parts.append(f"| {' | '.join(cells)} |")
        return "\n".join(parts)
    except ImportError:
        logger.warning("python-docx not installed. Attempting XML fallback.")
        return parse_docx_xml_fallback(file_bytes)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""

def parse_docx_xml_fallback(file_bytes: bytes) -> str:
    """Fallback XML parser for DOCX if docx module is missing."""
    try:
        file_like = io.BytesIO(file_bytes)
        with zipfile.ZipFile(file_like) as z:
            xml_content = z.read("word/document.xml")
            root = ET.fromstring(xml_content)
            # Match standard docx paragraph text tags
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_parts = []
            for t_el in root.findall('.//w:t', ns):
                if t_el.text:
                    text_parts.append(t_el.text)
            return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX XML fallback failed: {e}")
        return ""

def parse_pdf(file_bytes: bytes) -> str:
    """Parses PDF text contents using pypdf."""
    try:
        import pypdf
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n".join(parts)
    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        return ""


def parse_pdf_pages(file_bytes: bytes) -> List[ExtractedDocumentPage]:
    """Parses PDF text contents page-by-page using pypdf."""
    try:
        import pypdf
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append(ExtractedDocumentPage(page_number=index, text=text, source="pdf_text"))
        return pages
    except Exception as e:
        logger.error(f"PDF page parsing failed: {e}")
        return []


def ocr_image(file_bytes: bytes) -> str:
    """OCRs image bytes using Pillow and pytesseract when available."""
    try:
        from PIL import Image
        import pytesseract
        image = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(image) or ""
    except ImportError as e:
        raise RuntimeError("OCR dependencies are not installed. Install Pillow and pytesseract.") from e
    except Exception as e:
        raise RuntimeError(f"OCR image extraction failed: {e}") from e


def ocr_pdf_pages(file_bytes: bytes) -> List[ExtractedDocumentPage]:
    """OCRs scanned PDF pages when pdf2image, Pillow, pytesseract, and Poppler are available."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(file_bytes)
        pages = []
        for index, image in enumerate(images, start=1):
            text = pytesseract.image_to_string(image) or ""
            pages.append(ExtractedDocumentPage(page_number=index, text=text, source="ocr_pdf"))
        return pages
    except ImportError as e:
        raise RuntimeError("OCR PDF dependencies are not installed. Install pdf2image, Pillow, and pytesseract.") from e
    except Exception as e:
        raise RuntimeError(f"OCR PDF extraction failed: {e}") from e


def extract_document_pages(file_bytes: bytes, filename: str) -> DocumentExtractionResult:
    """
    Extracts document text with page metadata. Uses OCR for images and scanned PDFs
    when OCR dependencies and system binaries are available.
    """
    ext = os.path.splitext(filename)[1].lower()
    image_extensions = {".png", ".jpg", ".jpeg", ".tiff", ".tif"}

    if ext == ".pdf":
        pages = parse_pdf_pages(file_bytes)
        native_text = "\n\n".join(page.text for page in pages if page.text)
        if native_text.strip():
            return DocumentExtractionResult(
                text=native_text,
                pages=pages,
                extraction_method="pdf_text",
            )
        try:
            ocr_pages = ocr_pdf_pages(file_bytes)
            ocr_text = "\n\n".join(page.text for page in ocr_pages if page.text)
            return DocumentExtractionResult(
                text=ocr_text,
                pages=ocr_pages,
                extraction_method="ocr_pdf",
                ocr_status="completed" if ocr_text.strip() else "empty",
            )
        except RuntimeError as e:
            return DocumentExtractionResult(
                text="",
                pages=pages,
                extraction_method="pdf_text",
                ocr_status="unavailable",
                ocr_error=str(e),
            )

    if ext in image_extensions:
        try:
            text = ocr_image(file_bytes)
            return DocumentExtractionResult(
                text=text,
                pages=[ExtractedDocumentPage(page_number=1, text=text, source="ocr_image")],
                extraction_method="ocr_image",
                ocr_status="completed" if text.strip() else "empty",
            )
        except RuntimeError as e:
            return DocumentExtractionResult(
                text="",
                pages=[],
                extraction_method="ocr_image",
                ocr_status="unavailable",
                ocr_error=str(e),
            )

    text = extract_document_text(file_bytes, filename)
    return DocumentExtractionResult(
        text=text,
        pages=[ExtractedDocumentPage(page_number=1, text=text, source=ext.lstrip(".") or "text")],
        extraction_method=ext.lstrip(".") or "text",
    )

def extract_document_text(file_bytes: bytes, filename: str) -> str:
    """
    Unified entrypoint to parse and extract text from supported file types.
    """
    import os
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == ".pdf":
        return parse_pdf(file_bytes)
    elif ext == ".docx":
        return parse_docx(file_bytes)
    elif ext in (".txt", ".text"):
        return parse_txt_or_md(file_bytes)
    elif ext in (".md", ".markdown"):
        return parse_txt_or_md(file_bytes)
    elif ext == ".csv":
        return parse_csv(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return parse_xlsx(file_bytes)
        
    # Default fallback: try text decoding
    try:
        return file_bytes.decode("utf-8")
    except Exception:
        return ""
